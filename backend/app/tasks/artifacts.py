from __future__ import annotations

import logging
import os
import textwrap
from pathlib import Path

import httpx
import pdfplumber
from bs4 import BeautifulSoup
from docx import Document
from readability import Document as ReadabilityDocument
from sqlalchemy.orm import Session

from app.celery_app import celery_app
from app.core.database import SessionLocal
from app.models.artifact import AIArtifact, ArtifactStatus
from app.services import artifact_storage


logger = logging.getLogger(__name__)


def _with_db_session() -> Session:
    return SessionLocal()


@celery_app.task(name="artifacts.process_uploaded_artifact", max_retries=3)
def process_uploaded_artifact(artifact_id: int) -> None:
    db = _with_db_session()
    try:
        artifact = db.get(AIArtifact, artifact_id)
        if not artifact:
            logger.warning("Artifact %s not found", artifact_id)
            return
        if not artifact.s3_key:
            logger.error("Artifact %s missing s3_key", artifact_id)
            return
        tmp_path = artifact_storage.download_to_tempfile(artifact.s3_key)
        try:
            text = _extract_text(Path(tmp_path), artifact.source_details or {})
        finally:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass
        if not text:
            raise ValueError("Unable to extract text from document.")
        artifact.text_content = textwrap.shorten(text, 200000, placeholder=" …")
        artifact.status = ArtifactStatus.ready
        artifact.failure_reason = None
        db.commit()
    except Exception as exc:  # pylint: disable=broad-except
        logger.exception("Failed to process artifact %s", artifact_id)
        _mark_failed(db, artifact_id, str(exc))
    finally:
        db.close()


@celery_app.task(name="artifacts.scrape_job_description", max_retries=3)
def scrape_job_description(artifact_id: int) -> None:
    db = _with_db_session()
    try:
        artifact = db.get(AIArtifact, artifact_id)
        if not artifact:
            logger.warning("Artifact %s not found", artifact_id)
            return
        url = (artifact.source_details or {}).get("url")
        if not url:
            raise ValueError("Artifact missing URL.")
        text = _scrape_url(url)
        artifact.text_content = textwrap.shorten(text, 200000, placeholder=" …")
        artifact.status = ArtifactStatus.ready
        artifact.failure_reason = None
        db.commit()
    except Exception as exc:  # pylint: disable=broad-except
        logger.exception("Failed to scrape artifact %s", artifact_id)
        _mark_failed(db, artifact_id, str(exc))
    finally:
        db.close()


def _mark_failed(db: Session, artifact_id: int, reason: str) -> None:
    artifact = db.get(AIArtifact, artifact_id)
    if not artifact:
        return
    artifact.status = ArtifactStatus.failed
    artifact.failure_reason = reason[:500]
    db.commit()


def _extract_text(path: Path, details: dict) -> str:
    filename = (details or {}).get("filename", "").lower()
    if filename.endswith(".docx"):
        return _extract_docx(path)
    if filename.endswith(".pdf"):
        return _extract_pdf(path)
    return path.read_text(errors="ignore")


def _extract_docx(path: Path) -> str:
    document = Document(path)
    return "\n".join(p.text for p in document.paragraphs)


def _extract_pdf(path: Path) -> str:
    output: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            output.append(text)
    return "\n".join(output)


def _scrape_url(url: str) -> str:
    with httpx.Client(follow_redirects=True, timeout=30.0) as client:
        response = client.get(url)
        response.raise_for_status()
        raw_html = response.text
        readable = ReadabilityDocument(raw_html)
        summary_html = readable.summary(html_partial=True)
        soup = BeautifulSoup(summary_html, "html.parser")
        text = soup.get_text("\n")
        return text.strip()
